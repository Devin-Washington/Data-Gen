#!/usr/bin/env python3
"""
=============================================================================
OPERATION GROVE GUARDIAN — Synthetic Military Document Generator
=============================================================================
Generates cross-referenced military documents across operational phases for
the Republic of Solara FID/COIN scenario.

Documents per day (as applicable):
  - FRAGO  (daily)
  - ATO    (daily)
  - ACO    (daily)
  - JIPTL  (daily)
  - CCIR   (every CCIR_INTERVAL days + phase transitions)
  - PIR    (every PIR_INTERVAL days + phase transitions)
  - OPORD  (phase transitions only — 4 total)
  - ROE    (phase transitions + amendments)

Usage:
  pip install python-docx
  python ogg_generator.py                   # 8-day demo
  python ogg_generator.py --days 365        # full year
  python ogg_generator.py --days 100 --output ./my_output

Requirements: python-docx (pip install python-docx)
=============================================================================
"""

import os
import sys
import random
import argparse
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================================
# CONFIGURATION
# ============================================================================

# Phase definitions (day ranges)
PHASES = {
    1: {"name": "Shape",              "start": 0,   "end": 30,  "label": "I"},
    2: {"name": "Train and Secure",   "start": 31,  "end": 120, "label": "II"},
    3: {"name": "Operate and Expand", "start": 121, "end": 270, "label": "III"},
    4: {"name": "Transition",         "start": 271, "end": 999, "label": "IV"},
}

CCIR_INTERVAL = 3   # generate CCIR every N days
PIR_INTERVAL = 3    # generate PIR every N days

# Base date for D-Day
BASE_DATE = datetime(2026, 1, 20, 6, 0, 0)  # 200600ZJAN26

# Seed for reproducibility
random.seed(42)

# ============================================================================
# SCENARIO DATA POOLS — used to create realistic variation
# ============================================================================

SLM_LEADERS = [
    ("VIPER", "Comandante Elias Fuentes", "SLM Supreme Commander"),
    ("SCORPION", "Sub-Cmdr. Rafael Tovar", "SLM Military Chief"),
    ("GECKO", "Javier 'El Sombra' Mena", "SLM Corridor Cell Leader"),
    ("CORAL", "Lucia Varga", "SLM Propaganda Chief"),
    ("BUSHMASTER", "Carlos Delgado", "SLM Border Logistics Chief"),
    ("MOCCASIN", "Diego Soto", "SLM Swamp Sector Commander"),
    ("PYTHON", "Unknown True Name", "SLM Finance Cell Leader"),
    ("IGUANA", "Pablo Rios", "SLM Urban Cell Leader"),
    ("CAIMAN", "Hector Rivas", "SLM Maritime/Smuggling Coordinator"),
    ("ANACONDA", "Marco Gutierrez", "SLM Training Camp Commander"),
]

TARGET_POOL = [
    # (TGT_ID, Name, Category, Base MGRS, CDE, Phase_available_from)
    ("SLM-LOG-001", "Primary Weapons Cache (Border Region)", "LOGISTICS", "17R NM 4523 78", "LOW", 0),
    ("SLM-LOG-002", "Cross-Border Smuggling LZ Alpha", "LOGISTICS", "17R NM 5678 89", "LOW", 0),
    ("SLM-LOG-003", "Fuel/Supply Point (Scrubland East)", "LOGISTICS", "17R NM 4012 72", "LOW", 0),
    ("SLM-LOG-004", "IED/Demolitions Workshop", "LOGISTICS", "17R NM 3678 68", "MOD", 0),
    ("SLM-LOG-005", "Safe House Network (Company Town 3)", "LOGISTICS", "17R NM 3500 71", "HIGH", 0),
    ("SLM-LOG-006", "Secondary Weapons Cache (Swamp Edge)", "LOGISTICS", "17R NM 3100 65", "LOW", 15),
    ("SLM-LOG-007", "Smuggling LZ Bravo (Coastal Relay)", "LOGISTICS", "17R NM 5890 91", "LOW", 10),
    ("SLM-LOG-008", "Medical Supply Cache", "LOGISTICS", "17R NM 3345 69", "MOD", 30),
    ("SLM-LOG-009", "Vehicle Staging Area", "LOGISTICS", "17R NM 4200 74", "LOW", 45),
    ("SLM-LOG-010", "Boat Cache (River Junction)", "LOGISTICS", "17R NM 2800 62", "LOW", 20),
    ("SLM-C2-001", "Regional Commander CP (Swamp Central)", "C2", "17R NM 3456 67", "LOW", 0),
    ("SLM-C2-002", "Communications Relay (Hill 247)", "C2", "17R NM 3890 70", "LOW", 0),
    ("SLM-C2-003", "Finance/Funding Cell (Arcadia)", "C2", "17R NM 3210 65", "HIGH", 0),
    ("SLM-C2-004", "Alternate CP (Scrubland North)", "C2", "17R NM 3780 72", "LOW", 40),
    ("SLM-C2-005", "Courier Network Hub", "C2", "17R NM 3560 68", "MOD", 25),
    ("SLM-INF-001", "Propaganda Media Center (Arcadia Urban)", "INFO OPS", "17R NM 3200 65", "HIGH", 0),
    ("SLM-INF-002", "Social Media Operations Cell", "INFO OPS", "17R NM 3250 64", "HIGH", 0),
    ("SLM-INF-003", "Underground Print Shop (Company Town 7)", "INFO OPS", "17R NM 3400 66", "MOD", 20),
    ("SLM-TRN-001", "Training Camp Alpha (Deep Swamp)", "FORCE", "17R NM 2890 62", "LOW", 0),
    ("SLM-TRN-002", "Training Camp Bravo (Scrubland West)", "FORCE", "17R NM 2560 67", "LOW", 0),
    ("SLM-TRN-003", "Recruit Assembly Point (Lake Region)", "FORCE", "17R NM 2700 64", "LOW", 35),
    ("SLM-SAB-001", "Corridor Sabotage Cell (KM 30-50)", "FORCE", "17R NM 3900 75", "LOW", 5),
    ("SLM-SAB-002", "Corridor Sabotage Cell (KM 80-100)", "FORCE", "17R NM 4100 78", "LOW", 5),
    ("SLM-SAB-003", "Bridge Assault Team (Rio Verde Bridge)", "FORCE", "17R NM 4300 80", "MOD", 15),
    ("SLM-MAR-001", "Port Manatee Sleeper Cell", "FORCE", "17R NM 5500 88", "HIGH", 30),
    ("SLM-CYB-001", "Cyber Operations Cell (Unknown Location)", "C2", "17R NM 3200 65", "HIGH", 60),
]

MISSION_TYPES_ISR = [
    ("SHADOW", "MQ-1C", "UAS PLT/SOTF-K", "ISR"),
    ("RAVEN", "RQ-20B", "SOTF-C", "ISR"),
    ("PUMA", "RQ-11B", "CATF", "ISR"),
    ("SCAN EAGLE", "RQ-21A", "SOTF-K", "ISR"),
]

MISSION_TYPES_SUPPORT = [
    ("DUSTOFF", "UH-60M", "MEDEVAC DET", "MEDEVAC"),
    ("ATLAS", "C-146A", "AFSOC DET", "AIRLIFT"),
    ("TALON", "MC-130J", "AFSOC DET", "INFILTRATION"),
    ("NIGHTHAWK", "MH-60M", "160th SOAR DET", "ASSAULT SUPPORT"),
]

ISR_AREAS = [
    "CORRIDOR NORTH (COR-N)", "CORRIDOR SOUTH (COR-S)", "SWAMP CENTRAL (SWP-C)",
    "BORDER EAST (BDR-E)", "BORDER WEST (BDR-W)", "ARCADIA URBAN (ARC-U)",
    "PORT MANATEE APPROACH (PM-A)", "SCRUBLAND NORTH (SCR-N)", "SCRUBLAND SOUTH (SCR-S)",
    "LAKE REGION (LK-R)", "RIVER JUNCTION (RVR-J)", "COMPANY TOWN CLUSTER (CT-C)",
]

EVENTS_POOL = [
    # (day_trigger_min, day_trigger_max, event_text, impact_type)
    (1, 5, "SLM propaganda leaflets distributed in Company Town 2", "INFO"),
    (2, 8, "HNSF tip-line reports SLM movement near KM 45 of corridor", "INTEL"),
    (3, 10, "Small arms fire exchanged between HNSF patrol and SLM cell near the swamp edge", "CONTACT"),
    (5, 15, "IED discovered and neutralized on rail line at KM 67", "SABOTAGE"),
    (7, 20, "SLM propaganda video posted on social media showing staged attack", "INFO"),
    (8, 25, "HUMINT source reports SLM planning meeting in Scrubland East", "INTEL"),
    (10, 30, "Cross-border smuggling boat interdicted by HNSF Border Guard; weapons recovered", "INTERDICTION"),
    (12, 35, "CA team conducts well-digging project in Company Town 4; positive reception", "CIVIL"),
    (15, 40, "SLM ambush of CSG patrol near KM 92; 2x HNSF WIA, SLM repelled", "CONTACT"),
    (18, 50, "Radio Solara Libre begins broadcasting; initial listener metrics positive", "INFO"),
    (20, 55, "First SLM defector surrenders via amnesty program; debriefed by J-2", "INTEL"),
    (25, 60, "CUAS system detects and defeats hostile drone near Camp Citrus", "CUAS"),
    (30, 70, "Swamp Rangers conduct first independent patrol; no contact", "HNSF_PROGRESS"),
    (35, 80, "SLM attacks citrus processing plant in Sector 3; minimal damage", "SABOTAGE"),
    (40, 90, "Second amnesty defector provides SLM OOB information", "INTEL"),
    (45, 100, "CSG interdicts SLM sabotage cell preparing to attack Rio Verde Bridge", "INTERDICTION"),
    (50, 110, "SLM cyber intrusion detected on HNSF email server; isolated and eradicated", "CYBER"),
    (60, 130, "Swamp Rangers clear SLM camp; 12 SLM detained by HNSF", "CLEARING"),
    (75, 150, "SLM corridor attacks decrease 40% from baseline", "PROGRESS"),
    (90, 180, "GoS announces land reform pilot program for Grove Laborers", "CIVIL"),
    (100, 200, "SLM VIPER (supreme commander) communicates frustration to foreign sponsor per SIGINT", "INTEL"),
    (120, 240, "HNSF conducts first fully independent corridor security rotation", "HNSF_PROGRESS"),
    (150, 270, "SLM recruitment rate down estimated 60% per HUMINT", "PROGRESS"),
    (200, 300, "SLM offers ceasefire negotiations through intermediary", "DIPLOMATIC"),
    (250, 350, "GoS Ministry of Information assumes lead for Radio Solara Libre", "TRANSITION"),
]

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def get_phase(day):
    """Return phase number for a given operational day."""
    for phase_num, info in PHASES.items():
        if info["start"] <= day <= info["end"]:
            return phase_num
    return 4

def get_phase_info(day):
    """Return phase dict for a given day."""
    return PHASES[get_phase(day)]

def is_phase_transition(day):
    """Check if this day is the start of a new phase."""
    return any(info["start"] == day for info in PHASES.values())

def mil_dtg(dt):
    """Format datetime as military DTG."""
    return dt.strftime("%d%H%MZ%b%y").upper()

def day_date(day_offset):
    """Get the datetime for a given day offset from D-Day."""
    return BASE_DATE + timedelta(days=day_offset)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_classification_header_footer(doc, text="UNCLASSIFIED"):
    for section in doc.sections:
        hp = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        hp.text = text
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hp.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 128, 0)
        fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        fp.text = text
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in fp.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 128, 0)

def set_narrow_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

def add_heading_block(doc, lines, size=11):
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(size)

def add_para(doc, text, bold=False, size=12, indent=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent * 0.5)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def jitter_mgrs(base_mgrs):
    """Add small random variation to an MGRS coordinate string."""
    parts = base_mgrs.split()
    if len(parts) >= 4:
        try:
            e = int(parts[2]) + random.randint(-50, 50)
            n = int(parts[3]) + random.randint(-50, 50)
            return f"{parts[0]} {parts[1]} {e:04d} {n:02d}"
        except (ValueError, IndexError):
            pass
    return base_mgrs


# ============================================================================
# DAILY STATE ENGINE
# ============================================================================

class DailyState:
    """Tracks the evolving operational state for a given day."""

    def __init__(self, day):
        self.day = day
        self.date = day_date(day)
        self.dtg = mil_dtg(self.date)
        self.eff_dtg = mil_dtg(self.date)
        self.end_dtg = mil_dtg(self.date + timedelta(hours=18))
        self.phase = get_phase(day)
        self.phase_info = get_phase_info(day)
        self.phase_label = self.phase_info["label"]
        self.phase_name = self.phase_info["name"]

        # Evolving metrics
        self.slm_strength = max(500, 3000 - int(day * 6.5) + random.randint(-100, 100))
        self.hnsf_readiness = min(95, 35 + int(day * 0.18) + random.randint(-5, 5))
        self.corridor_threat = max(1, 10 - int(day * 0.02) + random.randint(-1, 1))  # 1-10
        self.popular_support_gos = min(80, 30 + int(day * 0.14) + random.randint(-3, 3))
        self.tip_line_calls = max(0, int(day * 0.8) + random.randint(-5, 10))
        self.amnesty_surrenders = max(0, int(day * 0.15) + random.randint(-2, 3))

        # Select today's events
        self.events = [e for e in EVENTS_POOL if e[0] <= day <= e[1]]
        if self.events:
            self.events = random.sample(self.events, min(len(self.events), random.randint(1, 3)))
        self.event_texts = [e[2] for e in self.events]

        # Select today's active targets
        available = [t for t in TARGET_POOL if t[5] <= day]
        # Some targets get "neutralized" over time
        neutralized_count = min(len(available) - 5, int(day * 0.03))
        random.seed(42 + day)  # deterministic per day
        if neutralized_count > 0:
            neutralized_ids = set(t[0] for t in random.sample(available, neutralized_count))
        else:
            neutralized_ids = set()
        self.active_targets = [t for t in available if t[0] not in neutralized_ids]
        random.shuffle(self.active_targets)

        # JIPTL: prioritize top 8-12, rest below cut line
        n_above_cut = min(len(self.active_targets), random.randint(6, 10))
        self.jiptl_above_cut = self.active_targets[:n_above_cut]
        self.jiptl_below_cut = self.active_targets[n_above_cut:]

        # ATO missions: generate ISR missions aligned to JIPTL target areas
        self.ato_missions = self._generate_ato_missions()

        # ACM variations
        self.acm_count = random.randint(6, 10)
        self.fscm_count = random.randint(5, 8)

        # Reset seed
        random.seed(42)

    def _generate_ato_missions(self):
        missions = []
        msn_num = 1

        # ISR missions covering JIPTL priority targets
        isr_areas_today = random.sample(ISR_AREAS, min(len(ISR_AREAS), random.randint(4, 7)))
        for area in isr_areas_today:
            platform = random.choice(MISSION_TYPES_ISR)
            callsign = f"{platform[0]} {random.randint(1, 99):02d}"
            start_hr = random.randint(6, 10)
            end_hr = start_hr + random.randint(4, 12)
            start_t = self.date + timedelta(hours=start_hr)
            end_t = self.date + timedelta(hours=min(end_hr, 23))
            missions.append({
                "msn_num": f"ATO-{msn_num:03d}",
                "callsign": callsign,
                "acft": platform[1],
                "unit": platform[2],
                "msn_type": platform[3],
                "target_area": area,
                "tot": f"{mil_dtg(start_t)}-{mil_dtg(end_t)}",
                "remarks": f"ISR coverage; ATO Day {self.day+1:03d}",
            })
            msn_num += 1

        # Support missions (MEDEVAC, airlift always present)
        for sup in MISSION_TYPES_SUPPORT[:2]:  # DUSTOFF + ATLAS always
            callsign = f"{sup[0]} {random.randint(1, 20):02d}"
            missions.append({
                "msn_num": f"ATO-{msn_num:03d}",
                "callsign": callsign,
                "acft": sup[1],
                "unit": sup[2],
                "msn_type": sup[3],
                "target_area": "CAMP CITRUS / JOA-WIDE",
                "tot": f"{mil_dtg(self.date)}-{self.end_dtg}",
                "remarks": "Standing mission" if sup[3] == "MEDEVAC" else "Resupply run",
            })
            msn_num += 1

        # Phase-dependent extra missions
        if self.phase >= 3:
            # More assault support in later phases
            extra = random.choice(MISSION_TYPES_SUPPORT[2:])
            callsign = f"{extra[0]} {random.randint(1, 20):02d}"
            missions.append({
                "msn_num": f"ATO-{msn_num:03d}",
                "callsign": callsign,
                "acft": extra[1],
                "unit": extra[2],
                "msn_type": extra[3],
                "target_area": random.choice(ISR_AREAS[:3]),
                "tot": f"{mil_dtg(self.date + timedelta(hours=random.randint(1,6)))}",
                "remarks": "HNSF-led operation support",
            })
            msn_num += 1

        return missions


# ============================================================================
# DOCUMENT GENERATORS
# ============================================================================

# ------- OPORD (Phase transitions only) -------

def generate_opord(state, output_dir):
    doc = Document()
    set_narrow_margins(doc)
    phase = state.phase
    phase_name = state.phase_name

    add_heading_block(doc, [
        "UNCLASSIFIED",
        "",
        f"Copy 1 of 15 Copies",
        "JOINT TASK FORCE - GROVE GUARDIAN",
        "Camp Citrus, Arcadia, Republic of Solara",
        state.dtg,
        "",
        f"OPORD {phase:03d}-26 (OPERATION GROVE GUARDIAN - PHASE {state.phase_label}: {phase_name.upper()}) (UNCLASSIFIED)",
    ])

    add_para(doc, "References:", bold=True)
    refs = [
        "a. Map, Series Z901, Republic of Solara, Sheets 1-4, Edition 3, 1:250,000.",
        "b. USSOCOM OPORD 25-007 (OPERATION SOUTHERN RESOLVE), DTG 150800ZDEC25.",
        "c. U.S. Embassy Solara Country Team Assessment, 01 Dec 2025.",
        "d. FM 5-0, Planning and Orders Production.",
        "e. JP 3-22, Foreign Internal Defense.",
        "f. JP 3-24, Counterinsurgency.",
    ]
    if phase > 1:
        refs.append(f"g. JTF-GG OPORD {phase-1:03d}-26 (Previous Phase).")
    for r in refs:
        add_para(doc, r, indent=1, space_after=2)

    add_para(doc, f"Time Zone: ZULU (Z)", bold=True)
    add_para(doc, "Task Organization: See Annex A (Task Organization).", bold=True)

    # Para 1 - Situation
    doc.add_heading("1. SITUATION", level=1)
    doc.add_heading("a. Area of Interest", level=2)
    add_para(doc, "The JOA encompasses the Republic of Solara and the border region extending 50km beyond the international boundary, including maritime approaches to Port Manatee. Refer to Annex B (Intelligence).")

    doc.add_heading("b. Assigned Area", level=2)
    add_para(doc, f"(1) Terrain. Republic of Solara: landlocked nation dominated by swampy watershed, pine scrub, cattle ranches, and citrus groves. Key terrain: Port Manatee corridor, capital Arcadia, citrus processing facilities, SLM swamp sanctuaries. Refer to Annex B.")
    add_para(doc, f"(2) Weather. {'Dry season; optimal conditions for ground ops and ISR.' if state.day < 150 else 'Wet season approaching; degraded ground mobility in swamp regions anticipated.'} Refer to Annex B.")

    doc.add_heading("c. Enemy Forces", level=2)
    add_para(doc, f"(1) The SLM is an irregular force currently estimated at {state.slm_strength} fighters in 15-20 cells. Recruited from Grove Laborers; sheltered by Scrub Folk. Funded by neighboring adversary.")
    if phase == 1:
        add_para(doc, "(2) SLM maintains initiative along the corridor with frequent sabotage. Swamp sanctuaries are largely uncontested. External support flowing freely across the border.")
    elif phase == 2:
        add_para(doc, "(2) SLM corridor attacks continuing but HNSF response improving. ISR providing increased early warning. SLM adapting tactics, shifting to nighttime operations. Border interdiction beginning to constrain supply.")
    elif phase == 3:
        add_para(doc, "(2) SLM under pressure across all operating areas. Corridor attacks reduced. Swamp sanctuaries being contested. Amnesty program generating defections. SLM leadership showing signs of internal friction.")
    else:
        add_para(doc, "(2) SLM significantly degraded. Remnants operating in small, isolated cells. Leadership fragmented. External support severely disrupted. SLM propaganda losing effectiveness.")

    doc.add_heading("d. Friendly Forces", level=2)
    add_para(doc, f"(1) Higher HQ Two Levels Up. USSOCOM.")
    add_para(doc, f"(2) Higher HQ One Level Up. USSOUTHCOM/TSOC.")
    add_para(doc, f"(3) HNSF readiness assessed at {state.hnsf_readiness}%. GoS popular support at approximately {state.popular_support_gos}%.")

    doc.add_heading("e. Civil Considerations", level=2)
    add_para(doc, "Population divided into Orchard Barons, Grove Laborers (center of gravity), and Scrub Folk. " +
             (f"Tip-line calls averaging {state.tip_line_calls}/day. {state.amnesty_surrenders} total amnesty surrenders to date." if state.day > 0 else "Initial engagement underway."))

    doc.add_heading("g. Assumptions", level=2)
    add_para(doc, "(1) GoS maintains political will. (2) HNSF remains loyal. (3) Neighboring adversary continues covert (not overt) support. (4) Port Manatee corridor remains sole economic export route.")

    # Para 2 - Mission
    doc.add_heading("2. MISSION", level=1)
    mission_verbs = {
        1: "establishes the JTF, conducts initial assessments, and begins advisory operations",
        2: "conducts intensive HNSF training and establishes corridor security",
        3: "conducts HNSF-led clearing operations and expands security",
        4: "transitions security lead to HNSF and prepares for redeployment",
    }
    add_para(doc, f"Effective {state.eff_dtg}, JTF-GROVE GUARDIAN {mission_verbs[phase]} by, with, and through Host Nation Security Forces in the Republic of Solara in order to neutralize the SLM, secure the Port Manatee economic corridor, and address root socio-economic grievances, setting conditions for a stable Republic of Solara.", bold=True)

    # Para 3 - Execution
    doc.add_heading("3. EXECUTION", level=1)
    doc.add_heading("a. Commander's Intent", level=2)
    intents = {
        1: "Establish advisory relationships and ISR architecture to set conditions for decisive operations in subsequent phases.",
        2: "Build HNSF capacity to independently secure the corridor while degrading SLM freedom of movement.",
        3: "Press the advantage. HNSF-led operations deny SLM sanctuary while MISO and CA efforts accelerate population support for the GoS.",
        4: "Ensure HNSF sustainability. Complete transition of all operations. Residual SOF capability for CT only.",
    }
    add_para(doc, f"Purpose: {intents[phase]}")

    main_efforts = {1: "LOE 1 (Develop HNSF)", 2: "LOE 2 (Secure Corridor)", 3: "LOE 3 (Counter SLM)", 4: "LOE 1 (Sustain HNSF)"}
    add_para(doc, f"Main Effort: {main_efforts[phase]}")
    add_para(doc, "End State: GoS and HNSF independently maintain security; corridor secure; SLM neutralized; governance reforms underway.")

    doc.add_heading("b. Concept of Operations", level=2)
    add_para(doc, f"Phase {state.phase_label} ({phase_name}) operations focus on {main_efforts[phase]}. Four LOEs remain mutually supporting. OPERATION RESOLUTE VOICE (MISO) supports all LOEs.")

    doc.add_heading("c. Tasks to Subordinate Units", level=2)
    tasks = {
        1: [
            ("SOTF-C", "Conduct initial assessment of Solaran Army units; begin Swamp Rangers training program."),
            ("SOTF-K", "Establish ISR architecture along Port Manatee corridor; begin CSG advisory operations."),
            ("SOTF-B", "Assess border security gaps; begin Border Guard training."),
            ("MISTF", "Establish Radio Solara Libre; begin MISO Phase I (Prepare and Legitimize)."),
            ("CATF", "Conduct initial civil-military engagement in priority company towns."),
        ],
        2: [
            ("SOTF-C", "Conduct intensive Swamp Rangers training; begin accompanied patrols into swamp periphery."),
            ("SOTF-K", "MAIN EFFORT. Establish layered corridor defense with CSG; achieve initial operating capability."),
            ("SOTF-B", "Conduct border interdiction training; execute joint patrols with HNSF Border Guard."),
            ("MISTF", "Execute MISO Phase II (Disrupt and Persuade); heavily market amnesty program."),
            ("CATF", "Scale civil-military projects; begin governance reform advocacy."),
        ],
        3: [
            ("SOTF-C", "MAIN EFFORT. Advise and accompany Swamp Rangers in clearing SLM sanctuaries."),
            ("SOTF-K", "Sustain and expand corridor security; achieve full operating capability."),
            ("SOTF-B", "Intensify border interdiction; dismantle SLM smuggling networks."),
            ("MISTF", "Intensify MISO targeting SLM leadership; begin Grey/Black operations."),
            ("CATF", "Expand CA projects to secondary company towns; advocate land reform."),
        ],
        4: [
            ("SOTF-C", "Transition Swamp Rangers to independent operations; maintain advisory presence."),
            ("SOTF-K", "Transfer corridor security lead to CSG; reduce advisory footprint."),
            ("SOTF-B", "Transition border operations to HNSF; prepare for redeployment."),
            ("MISTF", "Transfer MISO lead to GoS Ministry of Information."),
            ("CATF", "Transition CA programs to USAID and GoS agencies."),
        ],
    }
    for i, (unit, task) in enumerate(tasks[phase], 1):
        add_para(doc, f"({i}) {unit}. {task}", indent=1, space_after=4)

    doc.add_heading("d. Coordinating Instructions", level=2)
    add_para(doc, f"(1) This OPORD effective {state.eff_dtg}.")
    add_para(doc, "(2) CCIR: See separate CCIR document.")
    add_para(doc, "(3) ROE: See separate ROE document. U.S. forces under USSOUTHCOM ROE as supplemented.")
    add_para(doc, "(4) Fire Support/Airspace: See ACO and ATO.")

    # Para 4 - Sustainment
    doc.add_heading("4. SUSTAINMENT", level=1)
    add_para(doc, "Priority: (1) Main effort SOTF; (2) Other SOTFs; (3) MISTF; (4) CATF. Refer to Annex F.")
    add_para(doc, "a. Logistics. Resupply via air to Camp Citrus. Role 1 medical at Camp Citrus; MEDEVAC within 60 min.")

    # Para 5 - Command and Signal
    doc.add_heading("5. COMMAND AND SIGNAL", level=1)
    add_para(doc, "a. CDR JTF-GG at Camp Citrus. PACE: SATCOM / HF / Iridium / Runner.")
    add_para(doc, "b. Succession: CDR JTF-GG > DCDR > CDR SOTF-K > CDR SOTF-C.")

    add_para(doc, "")
    add_para(doc, "ACKNOWLEDGE:", bold=True)
    add_para(doc, "Acknowledgement means received and understood.")
    add_para(doc, "")
    add_para(doc, "J.R. MACKENZIE", bold=True)
    add_para(doc, "Major General, USA")
    add_para(doc, "Commanding")

    annexes = ["A-Task Organization","B-Intelligence","C-Operations","D-Fires","E-Protection",
               "F-Sustainment","G-Engineer","H-Signal","I-Air/Missile Defense","J-Public Affairs",
               "K-Civil Affairs","L-Information Collection","M-Assessment","N-Space Ops",
               "O-Omitted","P-Host-Nation Support","Q-KM","R-Reports","S-STO","T-Omitted",
               "U-IG","V-Interagency","W-OCS","X-Omitted","Y-Omitted","Z-Distribution"]
    add_para(doc, "ANNEXES:", bold=True)
    for a in annexes:
        add_para(doc, f"Annex {a}", indent=1, space_after=1)

    add_classification_header_footer(doc)
    fname = f"OPORD_{phase:03d}-26_Phase_{state.phase_label}_{phase_name.replace(' ','_')}.docx"
    path = os.path.join(output_dir, fname)
    doc.save(path)
    return path


# ------- FRAGO (Daily) -------

def generate_frago(state, output_dir, frago_num):
    doc = Document()
    set_narrow_margins(doc)

    add_heading_block(doc, [
        "UNCLASSIFIED",
        "",
        "JOINT TASK FORCE - GROVE GUARDIAN",
        f"FRAGMENTARY ORDER {frago_num:04d}-26",
        f"TO OPORD {state.phase:03d}-26 (PHASE {state.phase_label}: {state.phase_name.upper()})",
        state.dtg,
    ])

    add_para(doc, f"References: JTF-GG OPORD {state.phase:03d}-26; ATO {state.day+1:03d}-26; ACO {state.day+1:03d}-26.", bold=True)
    add_para(doc, f"Time Zone: ZULU (Z). ATO Day: {state.day+1:03d}.")

    doc.add_heading("1. SITUATION", level=1)
    add_para(doc, f"a. Enemy. SLM strength estimated at {state.slm_strength}. Corridor threat level: {state.corridor_threat}/10. HNSF readiness: {state.hnsf_readiness}%. GoS popular support: ~{state.popular_support_gos}%.")

    if state.event_texts:
        add_para(doc, "b. Significant Activities (last 24 hours):", bold=True)
        for evt in state.event_texts:
            add_para(doc, f"- {evt}", indent=1, space_after=2)
    else:
        add_para(doc, "b. No significant activities in the last 24 hours.")

    add_para(doc, f"c. Tip-line calls (24hr): {max(0, state.tip_line_calls + random.randint(-3,3))}. Cumulative amnesty surrenders: {state.amnesty_surrenders}.")

    doc.add_heading("2. MISSION", level=1)
    add_para(doc, "No change to OPORD mission statement.")

    doc.add_heading("3. EXECUTION", level=1)
    add_para(doc, f"a. Main Effort: No change. Phase {state.phase_label} ({state.phase_name}) operations continue.")

    add_para(doc, "b. Changes to tasks:", bold=True)
    # Generate 1-3 task changes per day
    task_changes = []
    if state.events:
        for evt in state.events:
            if evt[3] == "CONTACT":
                task_changes.append(f"SOTF-K: Increase CSG patrol frequency in sector of contact. Surge ISR to affected area.")
            elif evt[3] == "SABOTAGE":
                task_changes.append(f"SOTF-K: Coordinate with HNSF engineers for rapid repair. Adjust CSG patrol pattern to cover vulnerability.")
            elif evt[3] == "INTEL":
                task_changes.append(f"J-2: Develop intelligence lead. Nominate for ISR collection on next ATO cycle.")
            elif evt[3] == "INTERDICTION":
                task_changes.append(f"SOTF-B: Exploit captured material. Update border interdiction priorities.")
            elif evt[3] == "INFO":
                task_changes.append(f"MISTF: Develop counter-narrative. Coordinate with Radio Solara Libre for broadcast.")
            elif evt[3] == "CUAS":
                task_changes.append(f"J-3: Review CUAS posture at all FOBs. Report hostile UAS characteristics to J-2 for analysis.")
            elif evt[3] == "CYBER":
                task_changes.append(f"J-6/DCO: Elevate network monitoring. Implement recommended OPSEC changes.")
            elif evt[3] == "CLEARING":
                task_changes.append(f"SOTF-C: Exploit site for intelligence. Coordinate with J-2 for detainee processing via HNSF.")
            elif evt[3] == "CIVIL":
                task_changes.append(f"CATF: Continue engagement. Report atmospherics to J-2 and MISTF.")
            else:
                task_changes.append(f"All units: Continue Phase {state.phase_label} operations as directed.")

    if not task_changes:
        task_changes.append(f"No changes. Continue Phase {state.phase_label} operations IAW OPORD {state.phase:03d}-26.")

    for i, tc in enumerate(task_changes, 1):
        add_para(doc, f"({i}) {tc}", indent=1, space_after=3)

    add_para(doc, f"c. CCIR update: {'See updated CCIR document this period.' if state.day % CCIR_INTERVAL == 0 else 'No change to CCIR.'}")
    add_para(doc, f"d. JIPTL: See JIPTL {state.day+1:03d}-26 for current target priorities.")
    add_para(doc, f"e. ATO/ACO: See ATO {state.day+1:03d}-26 and ACO {state.day+1:03d}-26.")

    doc.add_heading("4. SUSTAINMENT", level=1)
    add_para(doc, "No change unless specified below.")
    if any(e[3] == "CONTACT" for e in state.events):
        add_para(doc, "- MEDEVAC: Confirm DUSTOFF status and blood product availability following contact.", indent=1)

    doc.add_heading("5. COMMAND AND SIGNAL", level=1)
    add_para(doc, "No change.")

    add_para(doc, "")
    add_para(doc, "ACKNOWLEDGE:", bold=True)
    add_para(doc, "For the Commander:")
    add_para(doc, "D.L. SANTOS, COL, USA")
    add_para(doc, "Chief of Staff, JTF-GROVE GUARDIAN")

    add_classification_header_footer(doc)
    fname = f"FRAGO_{frago_num:04d}-26_Day_{state.day+1:03d}.docx"
    path = os.path.join(output_dir, fname)
    doc.save(path)
    return path


# ------- ATO (Daily) -------

def generate_ato(state, output_dir):
    doc = Document()
    set_narrow_margins(doc)

    add_heading_block(doc, [
        "UNCLASSIFIED",
        "",
        "JOINT TASK FORCE - GROVE GUARDIAN",
        f"AIR TASKING ORDER {state.day+1:03d}-26",
        f"DTG: {state.dtg}",
        f"EFFECTIVE: {state.eff_dtg} TO {state.end_dtg}",
        f"ATO DAY: {state.day+1:03d} | PHASE {state.phase_label}: {state.phase_name.upper()}",
    ])

    doc.add_heading("SECTION 1: GENERAL", level=1)
    add_para(doc, f"References: OPORD {state.phase:03d}-26; ACO {state.day+1:03d}-26; ROE (current).")
    add_para(doc, f"Situation: Phase {state.phase_label} operations. SLM threat level along corridor: {state.corridor_threat}/10. HNSF readiness: {state.hnsf_readiness}%.")
    add_para(doc, f"Weather: {'Dry season; CAVU expected. Morning fog in swamp areas 0500-0800L.' if state.day < 150 else 'Wet season; scattered thunderstorms possible. Ceiling variable 800-unlimited.'}")

    doc.add_heading("SECTION 2: MISSION TASKING", level=1)
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    headers = ["MSN #", "CALLSIGN", "ACFT TYPE", "UNIT", "MSN TYPE", "TARGET/AREA", "TOT/ON-STATION", "REMARKS"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
        set_cell_shading(table.rows[0].cells[i], "D9E2F3")

    for m in state.ato_missions:
        row = table.add_row()
        vals = [m["msn_num"], m["callsign"], m["acft"], m["unit"], m["msn_type"], m["target_area"], m["tot"], m["remarks"]]
        for i, val in enumerate(vals):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    doc.add_heading("SECTION 3: SPINS (SUMMARY)", level=1)
    add_para(doc, "a. All UAS ops comply with ACO. Lost link: RTB profile. Notify JTF JOC immediately.")
    add_para(doc, "b. UAS will not overfly populated areas below 500ft AGL without JTF JOC approval.")
    add_para(doc, f"c. MEDEVAC: DUSTOFF on standby; 60-min response. FM 38.50 / SATCOM GROVE-MED-1.")
    add_para(doc, "d. CUAS: Report hostile UAS to JTF JOC. EW engagement authorized; kinetic requires CDR approval except self-defense.")
    add_para(doc, "e. ROE: Per current ROE. PID required for all engagements.")
    add_para(doc, f"f. Total missions this ATO: {len(state.ato_missions)}. ISR: {sum(1 for m in state.ato_missions if m['msn_type']=='ISR')}. Support: {sum(1 for m in state.ato_missions if m['msn_type']!='ISR')}.")

    add_classification_header_footer(doc)
    fname = f"ATO_{state.day+1:03d}-26_Day_{state.day+1:03d}.docx"
    path = os.path.join(output_dir, fname)
    doc.save(path)
    return path


# ------- ACO (Daily) -------

def generate_aco(state, output_dir):
    doc = Document()
    set_narrow_margins(doc)

    add_heading_block(doc, [
        "UNCLASSIFIED",
        "",
        "JOINT TASK FORCE - GROVE GUARDIAN",
        f"AIRSPACE CONTROL ORDER {state.day+1:03d}-26",
        f"DTG: {state.dtg}",
        f"EFFECTIVE: {state.eff_dtg} TO {state.end_dtg}",
    ])

    doc.add_heading("1. GENERAL", level=1)
    add_para(doc, f"ACMs approved by ACA (JTF J-3/Air) for ATO Day {state.day+1:03d}. Disseminated via this ACO.")

    doc.add_heading("2. AIRSPACE COORDINATING MEASURES (ACMs)", level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    acm_h = ["ACM #", "TYPE", "NAME", "LOCATION", "ALTITUDE", "EFFECTIVE", "CTRL AGENCY"]
    for i, h in enumerate(acm_h):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
        set_cell_shading(table.rows[0].cells[i], "D9E2F3")

    acm_types = ["ROZ", "UA", "HIDACZ", "ACA"]
    acm_names = ["SHADOW NORTH", "SHADOW SOUTH", "RAVEN CENTRAL", "RAVEN BORDER",
                 "SCAN PORT", "CORRIDOR SHIELD", "CAMP CITRUS", "PUMA URBAN",
                 "SWAMP OVERWATCH", "BORDER WATCH"]
    acm_bases = [
        ("N27 30 00 W081 45 00 / 20nm", "SFC-15000ft MSL"),
        ("N27 00 00 W081 30 00 / 20nm", "SFC-15000ft MSL"),
        ("N27 15 00 W081 50 00 / 10nm", "SFC-1200ft AGL"),
        ("N27 20 00 W081 15 00 / 10nm", "SFC-1200ft AGL"),
        ("N27 40 00 W082 30 00 / 15nm", "SFC-10000ft MSL"),
        ("N27 00-40 W081 30-W082 30 / 5nm corridor", "SFC-5000ft AGL"),
        ("N27 12 00 W081 46 00 / 3nm", "SFC-3000ft AGL"),
        ("N27 12 00 W081 46 00 / 5nm", "SFC-500ft AGL"),
        ("N27 10 00 W081 55 00 / 12nm", "SFC-8000ft MSL"),
        ("N27 25 00 W081 10 00 / 10nm", "SFC-5000ft AGL"),
    ]

    random.seed(42 + state.day)
    n_acms = min(state.acm_count, len(acm_names))
    selected = random.sample(range(len(acm_names)), n_acms)

    for idx, sel in enumerate(selected, 1):
        row = table.add_row()
        acm_type = acm_types[sel % len(acm_types)]
        vals = [
            f"ACM-{idx:02d}", acm_type, acm_names[sel],
            acm_bases[sel][0], acm_bases[sel][1],
            f"{state.eff_dtg}-{state.end_dtg}" if acm_type != "ACA" else "CONTINUOUS",
            "JTF JOC" if acm_type in ["ROZ","HIDACZ","ACA"] else random.choice(["SOTF-C","SOTF-K","SOTF-B","CATF"]),
        ]
        for i, val in enumerate(vals):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    doc.add_heading("3. FIRE SUPPORT COORDINATION MEASURES (FSCMs)", level=1)
    table2 = doc.add_table(rows=1, cols=6)
    table2.style = 'Table Grid'
    fscm_h = ["FSCM #", "TYPE", "NAME", "LOCATION", "EFFECTIVE", "EST. AUTH."]
    for i, h in enumerate(fscm_h):
        table2.rows[0].cells[i].text = h
        for p in table2.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
        set_cell_shading(table2.rows[0].cells[i], "E2EFDA")

    fscm_data = [
        ("NFA", "ARCADIA CITY", "5km radius Arcadia center", "CONTINUOUS"),
        ("NFA", "PORT MANATEE", "3km radius Port Manatee", "CONTINUOUS"),
        ("RFA", "CORRIDOR ZONE", "5km either side PM corridor", "CONTINUOUS"),
        ("NFA", "EMBASSY COMPOUND", "1km radius U.S. Embassy", "CONTINUOUS"),
        ("CFL", "BORDER CFL", "Along international border", "CONTINUOUS"),
        ("NFA", "HOSPITAL ZONE", "500m radius Arcadia General", "CONTINUOUS"),
        ("FFA", "SWAMP CLEAR", "Designated SLM sanctuary area", "ON ORDER"),
        ("RFA", "COMPANY TOWN BUFFER", "2km radius Company Towns 1-5", "CONTINUOUS"),
    ]
    n_fscms = min(state.fscm_count, len(fscm_data))
    for idx in range(n_fscms):
        row = table2.add_row()
        fd = fscm_data[idx]
        vals = [f"FSCM-{idx+1:02d}", fd[0], fd[1], fd[2], fd[3], "JTF CDR"]
        for i, val in enumerate(vals):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    add_para(doc, "")
    add_para(doc, f"Approved: R.P. THORNTON, COL, USAF — ACA (Delegated)", bold=True)

    random.seed(42)
    add_classification_header_footer(doc)
    fname = f"ACO_{state.day+1:03d}-26_Day_{state.day+1:03d}.docx"
    path = os.path.join(output_dir, fname)
    doc.save(path)
    return path


# ------- JIPTL (Daily) -------

def generate_jiptl(state, output_dir):
    doc = Document()
    set_narrow_margins(doc)
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        w, h = section.page_height, section.page_width
        section.page_width = w
        section.page_height = h

    add_heading_block(doc, [
        "UNCLASSIFIED",
        "",
        "JTF-GROVE GUARDIAN — JIPTL",
        f"JIPTL {state.day+1:03d}-26 | DTG: {state.dtg} | ATO DAY {state.day+1:03d}",
        f"PHASE {state.phase_label}: {state.phase_name.upper()}",
        "APPROVED: JTCB / JTF COMMANDER",
    ])

    add_para(doc, f"Targeting guidance: Priority (1) SLM logistics/external support; (2) SLM C2; (3) SLM sanctuaries; (4) SLM info ops. CDE required. HNSF concurrence required. Non-lethal preferred.", size=10)

    cols = 10
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    headers = ["PRI", "TGT ID", "TARGET NAME", "CAT", "LOCATION", "DESIRED EFFECT", "CDE", "NOMINATOR", "OBJ", "STATUS"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(7)
        set_cell_shading(table.rows[0].cells[i], "D9E2F3")

    effects = ["DESTROY", "NEUTRALIZE", "DENY", "DISRUPT", "DISRUPT (NON-LETHAL)", "DEGRADE"]
    nominators = ["SOTF-C", "SOTF-K", "SOTF-B", "MISTF", "CATF", "J-2"]
    objectives = ["OBJ 1: Secure Corridor", "OBJ 2: Neutralize SLM", "OBJ 3: Counter SLM", "OBJ 4: Isolate"]

    random.seed(42 + state.day)
    for pri, tgt in enumerate(state.jiptl_above_cut, 1):
        row = table.add_row()
        vals = [
            str(pri), tgt[0], tgt[1], tgt[2], jitter_mgrs(tgt[3]),
            random.choice(effects[:4]) if tgt[2] != "INFO OPS" else "DISRUPT (NON-LETHAL)",
            tgt[4],
            random.choice(nominators),
            random.choice(objectives),
            "NOMINATED",
        ]
        for i, val in enumerate(vals):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(7)

    # Cut line
    if state.jiptl_below_cut:
        cut_row = table.add_row()
        merged = cut_row.cells[0].merge(cut_row.cells[cols-1])
        merged.text = "--- CUT LINE ---"
        for p in merged.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(255, 0, 0)
        set_cell_shading(merged, "FFC7CE")

        for pri_offset, tgt in enumerate(state.jiptl_below_cut, len(state.jiptl_above_cut)+1):
            row = table.add_row()
            vals = [
                str(pri_offset), tgt[0], tgt[1], tgt[2], jitter_mgrs(tgt[3]),
                random.choice(effects),
                tgt[4],
                random.choice(nominators),
                random.choice(objectives),
                "BELOW CUT",
            ]
            for i, val in enumerate(vals):
                row.cells[i].text = val
                for p in row.cells[i].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(7)
            for cell in row.cells:
                set_cell_shading(cell, "FFF2CC")

    random.seed(42)
    add_classification_header_footer(doc)
    fname = f"JIPTL_{state.day+1:03d}-26_Day_{state.day+1:03d}.docx"
    path = os.path.join(output_dir, fname)
    doc.save(path)
    return path


# ------- ROE (Phase transitions + amendments) -------

def generate_roe(state, output_dir, roe_version=1):
    doc = Document()
    set_narrow_margins(doc)

    add_heading_block(doc, [
        "UNCLASSIFIED",
        "",
        "JOINT TASK FORCE - GROVE GUARDIAN",
        f"RULES OF ENGAGEMENT (VERSION {roe_version})",
        f"DTG: {state.dtg}",
        f"EFFECTIVE PHASE {state.phase_label}: {state.phase_name.upper()}",
    ])

    doc.add_heading("1. SITUATION", level=1)
    add_para(doc, f"JTF-GG conducts FID/COIN in the Republic of Solara. Phase {state.phase_label} ({state.phase_name}). SLM est. strength: {state.slm_strength}. These ROE supplement USSOUTHCOM Standing ROE.")

    doc.add_heading("2. GENERAL ROE", level=1)
    rules = [
        "2.1 Self-Defense: Inherent right retained at all times.",
        "2.2 Defense of HNSF: Authorized when HNSF subject to hostile act/intent and unable to self-defend. Authority: On-scene CDR (O-5+) or SOTF CDR.",
        "2.3 Defense of Designated Persons: Authorized per SJA-maintained designated persons list.",
        "2.4 Proportionality: Force proportional to threat; minimum necessary.",
        "2.5 Discrimination: PID required before engagement.",
        "2.6 Minimum Force: Use least force necessary. EOF procedures when time permits.",
    ]
    for r in rules:
        add_para(doc, r, space_after=4)

    doc.add_heading("3. SPECIFIC PROVISIONS", level=1)
    add_para(doc, "3.1 SLM Status: NOT declared hostile force. Engagement requires hostile act/hostile intent only.")
    add_para(doc, "3.2 Detention: NOT authorized. Transfer to HNSF immediately.")
    add_para(doc, "3.3 Sensitive Sites: 500m standoff without JTF CDR approval.")
    add_para(doc, "3.4 Cross-Border: NOT authorized without JTF CDR + USSOUTHCOM approval.")
    add_para(doc, "3.5 UAS: ISR only (unarmed). EW CUAS authorized; kinetic CUAS requires CDR approval except self-defense.")

    # Phase-specific amendments
    if state.phase >= 2:
        doc.add_heading("4. PHASE-SPECIFIC AMENDMENTS", level=1)
        if state.phase == 2:
            add_para(doc, "4.1 Accompanied Operations: U.S. advisors accompanying HNSF on corridor patrols may use force in collective self-defense of the combined element.")
            add_para(doc, "4.2 Escalation: During HNSF-led checkpoint operations, U.S. advisors will defer to HNSF EOF procedures unless U.S. lives are directly threatened.")
        elif state.phase == 3:
            add_para(doc, "4.1 Clearing Operations: During HNSF-led clearing operations in designated areas, U.S. advisors may call for fire support in coordination with HNSF commander when combined element is decisively engaged.")
            add_para(doc, "4.2 FFA Activation: JTF CDR may activate FFA SWAMP CLEAR (FSCM-07) for specific HNSF-led clearing operations. All fires require HNSF commander approval.")
        elif state.phase == 4:
            add_para(doc, "4.1 Reduced Posture: As HNSF assumes security lead, U.S. force protection posture remains unchanged. No reduction in self-defense authorities.")
            add_para(doc, "4.2 Residual CT: Any CT operations require JTF CDR + USSOUTHCOM approval.")

    doc.add_heading("5. EOF PROCEDURES", level=1)
    add_para(doc, "SHOUT > SHOW > SHOVE > SHOOT (Warning) > SHOOT (Neutralize)")

    doc.add_heading("6. REPORTING", level=1)
    add_para(doc, "All uses of force: Immediate report to JTF JOC. SIGACT within 1hr. Written statement within 24hr. SJA review within 72hr.")

    add_para(doc, "")
    add_para(doc, "J.R. MACKENZIE, MG, USA — Commander, JTF-GG", bold=True)
    add_para(doc, "Legal Review: T.M. HARGROVE, COL, JA — SJA, JTF-GG")

    add_classification_header_footer(doc)
    fname = f"ROE_V{roe_version:02d}_Phase_{state.phase_label}_{state.phase_name.replace(' ','_')}.docx"
    path = os.path.join(output_dir, fname)
    doc.save(path)
    return path


# ------- CCIR (Periodic) -------

def generate_ccir(state, output_dir, ccir_num):
    doc = Document()
    set_narrow_margins(doc)

    add_heading_block(doc, [
        "UNCLASSIFIED",
        "",
        "JTF-GROVE GUARDIAN",
        f"CCIR UPDATE {ccir_num:03d} | DTG: {state.dtg}",
        f"PHASE {state.phase_label}: {state.phase_name.upper()} | DAY {state.day+1:03d}",
    ])

    doc.add_heading("PRIORITY INTELLIGENCE REQUIREMENTS (PIR)", level=1)
    pir_table = doc.add_table(rows=1, cols=3)
    pir_table.style = 'Table Grid'
    for i, h in enumerate(["PIR", "REQUIREMENT", "DECISION POINT"]):
        pir_table.rows[0].cells[i].text = h
        for p in pir_table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(pir_table.rows[0].cells[i], "D9E2F3")

    # PIRs evolve by phase
    phase_pirs = {
        1: [
            ("PIR 1", "SLM intentions/timeline for corridor attacks", "Adjust CSG posture"),
            ("PIR 2", "SLM OOB, cell structure, leadership", "Prioritize SOTF-C ops"),
            ("PIR 3", "Cross-border logistics routes/methods", "Adjust SOTF-B interdiction"),
            ("PIR 4", "Popular support levels (Grove Laborers/Scrub Folk)", "Shape MISO strategy"),
            ("PIR 5", "SLM advanced weapons (MANPADS, UAS)", "Adjust force protection"),
            ("PIR 6", "SLM/sponsor cyber capabilities", "Elevate DCO posture"),
        ],
        2: [
            ("PIR 1", "SLM adaptation to corridor security measures", "Adjust CSG TTPs"),
            ("PIR 2", "SLM OOB changes and leadership movements", "Sequence clearing ops"),
            ("PIR 3", "Cross-border logistics; effectiveness of interdiction", "Adjust border ops"),
            ("PIR 4", "Population response to MISO and CA programs", "Redirect MISO/CA"),
            ("PIR 5", "SLM advanced weapons acquisition", "Adjust force protection"),
            ("PIR 6", "SLM cyber targeting of HNSF C2 systems", "Harden networks"),
        ],
        3: [
            ("PIR 1", "SLM intentions to escalate or negotiate", "Inform CDR decision on operational tempo"),
            ("PIR 2", "SLM remaining capability and cohesion", "Prioritize remaining clearing ops"),
            ("PIR 3", "Foreign sponsor commitment level", "Inform diplomatic approach"),
            ("PIR 4", "Population confidence in GoS reforms", "Advise GoS on reform pace"),
            ("PIR 5", "SLM IED/sabotage residual capability", "Maintain corridor security"),
            ("PIR 6", "SLM information ops effectiveness", "Counter remaining propaganda"),
        ],
        4: [
            ("PIR 1", "SLM reconstitution potential", "Determine residual CT requirement"),
            ("PIR 2", "HNSF sustainability without U.S. support", "Determine transition timeline"),
            ("PIR 3", "Foreign sponsor future intentions", "Inform post-transition posture"),
            ("PIR 4", "Population confidence trajectory", "Advise GoS on long-term stability"),
            ("PIR 5", "Residual SLM cells and spoiler potential", "Maintain awareness"),
            ("PIR 6", "GoS governance reform follow-through", "Shape final advisory messaging"),
        ],
    }
    for pir_data in phase_pirs.get(state.phase, phase_pirs[1]):
        row = pir_table.add_row()
        for i, val in enumerate(pir_data):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_heading("FRIENDLY FORCE INFORMATION REQUIREMENTS (FFIR)", level=1)
    ffir_table = doc.add_table(rows=1, cols=3)
    ffir_table.style = 'Table Grid'
    for i, h in enumerate(["FFIR", "REQUIREMENT", "DECISION POINT"]):
        ffir_table.rows[0].cells[i].text = h
        for p in ffir_table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(ffir_table.rows[0].cells[i], "E2EFDA")

    ffirs = [
        ("FFIR 1", "Loss of comms with any advisory team > 30 min", "Initiate PR procedures"),
        ("FFIR 2", "U.S. KIA/WIA/MIA", "MEDEVAC/PR; notify higher"),
        ("FFIR 3", f"HNSF readiness below 70% (current: {state.hnsf_readiness}%)", "Reassess advisory priorities"),
        ("FFIR 4", "HNSF disloyalty or refusal to operate", "Reassess partner vetting"),
        ("FFIR 5", "Corridor physical disruption", "Initiate rapid repair"),
        ("FFIR 6", "ROE violation or unauthorized use of force", "Initiate investigation"),
    ]
    for f in ffirs:
        row = ffir_table.add_row()
        for i, val in enumerate(f):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_heading("EEFI", level=1)
    eefis = ["SOF team locations/patterns", "ISR capabilities/gaps", "Intel sharing arrangements",
             "CUAS capabilities", "MEDEVAC/PR procedures", "Comms architecture", "HNSF op timelines"]
    for i, e in enumerate(eefis, 1):
        add_para(doc, f"EEFI {i}: {e}", indent=1, space_after=2)

    add_para(doc, "")
    add_para(doc, f"Next CCIR review: {mil_dtg(state.date + timedelta(days=CCIR_INTERVAL))}", bold=True)
    add_para(doc, "J.R. MACKENZIE, MG, USA — Commander, JTF-GG")

    add_classification_header_footer(doc)
    fname = f"CCIR_{ccir_num:03d}_Day_{state.day+1:03d}.docx"
    path = os.path.join(output_dir, fname)
    doc.save(path)
    return path


# ------- PIR (Periodic) -------

def generate_pir(state, output_dir, pir_num):
    doc = Document()
    set_narrow_margins(doc)

    add_heading_block(doc, [
        "UNCLASSIFIED",
        "",
        "JTF-GROVE GUARDIAN",
        f"PRIORITY INTELLIGENCE REQUIREMENTS UPDATE {pir_num:03d}",
        f"DTG: {state.dtg} | DAY {state.day+1:03d} | PHASE {state.phase_label}",
    ])

    # Detailed PIRs with indicators - evolve by phase
    pir_configs = {
        1: [
            ("PIR 1: SLM Corridor Attack Intentions",
             ["Increased comms along corridor", "Personnel/equipment movement toward infrastructure", "IED material pre-positioning", "Recon of CSG patrols"],
             "SIGINT, HUMINT (tip-line), UAS ISR (SHADOW)", "Continuous; 24hr cycle"),
            ("PIR 2: SLM Order of Battle",
             ["New cell identification", "Strength/org changes", "Training activities", "Leadership movement", "Recruitment in company towns"],
             "HUMINT (informants, defectors), SIGINT, UAS ISR, OSINT", "72hr update; immediate for leadership"),
            ("PIR 3: Cross-Border Logistics",
             ["Boat movement along coast", "Vehicle/pack movement at border", "New cache sites", "Financial transactions", "New weapons in SLM inventory"],
             "SIGINT, GEOINT, UAS ISR, HUMINT (border), liaison intel", "72hr prep; immediate for active ops"),
        ],
        2: [
            ("PIR 1: SLM Tactical Adaptation",
             ["Changes to attack TTPs", "Shift to nighttime ops", "Use of new IED types", "Counter-ISR measures", "Targeting of HNSF leadership"],
             "SIGINT, HUMINT, UAS ISR, TECHINT (recovered IEDs)", "Continuous; 24hr cycle"),
            ("PIR 2: SLM Leadership and Cohesion",
             ["Leadership disputes per SIGINT", "Cell fragmentation", "Defections/amnesty surrenders", "Changes in SLM messaging tone"],
             "SIGINT, HUMINT (defectors), OSINT (social media)", "48hr update"),
            ("PIR 3: Interdiction Effectiveness",
             ["Reduction in border crossings", "SLM supply shortages", "Foreign sponsor response to interdiction", "New smuggling routes"],
             "SIGINT, HUMINT, GEOINT, UAS ISR, liaison intel", "Weekly assessment"),
        ],
        3: [
            ("PIR 1: SLM Escalation vs Negotiation Intent",
             ["SLM leadership communications re: ceasefire", "Escalation of attacks as desperation", "Outreach to intermediaries", "Foreign sponsor guidance to SLM"],
             "SIGINT, HUMINT, diplomatic channels", "Immediate; 24hr cycle"),
            ("PIR 2: SLM Remaining Capability",
             ["Functional cells remaining", "Weapons/ammo stockpile status", "Recruit pipeline", "Morale indicators"],
             "All sources; defector debriefs critical", "48hr update"),
            ("PIR 3: Population Sentiment Trajectory",
             [f"Tip-line volume (current: ~{state.tip_line_calls}/day)", f"Amnesty rate (current: {state.amnesty_surrenders} total)", "GoS program participation", "Social media sentiment"],
             "HUMINT (CA), OSINT, polling", "Weekly"),
        ],
        4: [
            ("PIR 1: SLM Reconstitution Potential",
             ["Remaining leadership at large", "Foreign sponsor willingness to re-arm", "Recruitment potential", "Residual safe havens"],
             "All sources", "Weekly assessment"),
            ("PIR 2: HNSF Sustainability",
             [f"HNSF independent capability (current: {state.hnsf_readiness}%)", "Leadership quality", "Logistics self-sufficiency", "Intel collection capacity"],
             "Advisory team assessments, HNSF reporting", "Weekly"),
            ("PIR 3: GoS Reform Follow-Through",
             ["Land reform implementation", "Labor law enforcement", "GoS spending on company town services", "Orchard Baron cooperation"],
             "CA teams, Embassy reporting, OSINT", "Bi-weekly"),
        ],
    }

    for title, indicators, collection, ltiov in pir_configs.get(state.phase, pir_configs[1]):
        doc.add_heading(title, level=2)
        add_para(doc, "Indicators:", bold=True)
        for ind in indicators:
            add_para(doc, f"- {ind}", indent=1, space_after=2)
        add_para(doc, f"Collection: {collection}")
        add_para(doc, f"LTIOV: {ltiov}")
        add_para(doc, "")

    add_para(doc, f"Next PIR review: {mil_dtg(state.date + timedelta(days=PIR_INTERVAL))}", bold=True)
    add_para(doc, "S.K. NORTON, COL, MI — JTF J-2")

    add_classification_header_footer(doc)
    fname = f"PIR_{pir_num:03d}_Day_{state.day+1:03d}.docx"
    path = os.path.join(output_dir, fname)
    doc.save(path)
    return path


# ============================================================================
# MAIN ORCHESTRATOR
# ============================================================================

def run(num_days, output_root):
    """Generate all documents for the specified number of operational days."""

    # Create output directories
    dirs = {
        "OPORD": os.path.join(output_root, "OPORD"),
        "FRAGO": os.path.join(output_root, "FRAGO"),
        "ATO":   os.path.join(output_root, "ATO"),
        "ROE":   os.path.join(output_root, "ROE"),
        "ACO":   os.path.join(output_root, "ACO"),
        "CCIR":  os.path.join(output_root, "CCIR"),
        "PIR":   os.path.join(output_root, "PIR"),
        "JIPTL": os.path.join(output_root, "JIPTL"),
    }
    for d in dirs.values():
        os.makedirs(d, exist_ok=True)

    frago_counter = 0
    ccir_counter = 0
    pir_counter = 0
    roe_version = 0
    last_phase = 0
    total_docs = 0

    print(f"\n{'='*60}")
    print(f"OPERATION GROVE GUARDIAN — Document Generator")
    print(f"Generating {num_days} operational days")
    print(f"Output: {output_root}")
    print(f"{'='*60}\n")

    for day in range(num_days):
        state = DailyState(day)
        current_phase = state.phase
        phase_changed = (current_phase != last_phase)

        print(f"  Day {day+1:03d} (D+{day}) | Phase {state.phase_label}: {state.phase_name} | "
              f"SLM: {state.slm_strength} | HNSF: {state.hnsf_readiness}%", end="")

        day_docs = 0

        # OPORD — phase transitions only
        if phase_changed:
            generate_opord(state, dirs["OPORD"])
            day_docs += 1

        # ROE — phase transitions
        if phase_changed:
            roe_version += 1
            generate_roe(state, dirs["ROE"], roe_version)
            day_docs += 1

        # FRAGO — daily
        frago_counter += 1
        generate_frago(state, dirs["FRAGO"], frago_counter)
        day_docs += 1

        # ATO — daily
        generate_ato(state, dirs["ATO"])
        day_docs += 1

        # ACO — daily
        generate_aco(state, dirs["ACO"])
        day_docs += 1

        # JIPTL — daily
        generate_jiptl(state, dirs["JIPTL"])
        day_docs += 1

        # CCIR — every N days + phase transitions
        if day % CCIR_INTERVAL == 0 or phase_changed:
            ccir_counter += 1
            generate_ccir(state, dirs["CCIR"], ccir_counter)
            day_docs += 1

        # PIR — every N days + phase transitions
        if day % PIR_INTERVAL == 0 or phase_changed:
            pir_counter += 1
            generate_pir(state, dirs["PIR"], pir_counter)
            day_docs += 1

        total_docs += day_docs
        last_phase = current_phase
        print(f" | {day_docs} docs")

    print(f"\n{'='*60}")
    print(f"COMPLETE: {total_docs} documents generated across {num_days} days.")
    print(f"{'='*60}")
    print(f"\nOutput structure:")
    for name, path in dirs.items():
        count = len([f for f in os.listdir(path) if f.endswith('.docx')])
        print(f"  {name:8s}/ — {count} documents")
    print()


# ============================================================================
# CLI
# ============================================================================

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="OPERATION GROVE GUARDIAN Synthetic Data Generator")
    parser.add_argument("--days", type=int, default=8, help="Number of operational days to generate (default: 8)")
    parser.add_argument("--output", type=str, default="./OGG_Output", help="Output root directory (default: ./OGG_Output)")
    args = parser.parse_args()

    run(args.days, args.output)
